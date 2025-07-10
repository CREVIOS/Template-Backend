import io
import re
from datetime import datetime
from typing import List, Dict, Any, Union, Optional
from loguru import logger
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import json
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.oxml.shared import qn

logger.add("logs/document_exporter.log", rotation="100 MB", retention="7 days", level="INFO")

WHITELIST_TAGS = {
    'div', 'span', 'p', 'br', 'strong', 'em', 'u', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
    'ul', 'ol', 'li', 'table', 'tr', 'td', 'th', 'thead', 'tbody'
}


class ExportStyleConfig:
    """Runtime‑configurable styling and behaviour flags."""

    numbering_style: str = "numeric"  # numeric | roman | alpha
    drafting_note_as_comment: bool = True  # Set to True to test comments
    logo_path: Optional[str] = None  # path to footer logo

    def __init__(self, **kwargs):
        for k, v in kwargs.items():
            setattr(self, k, v)


class DocumentExporter:
    """Handle document export and preview generation with enhanced formatting using native DOCX XML"""

    PLACEHOLDER_REGEX = re.compile(r"(\[[^\]]+\])")

    def __init__(self):
        self.logger = logger
        self.numbering_part = None
        self.abstract_num_id = 0
        self.num_id = 1
        self.comment_counter = 0

    def export_json_to_docx(
        self,
        payload: Union[str, Dict[str, Any]],
        config: Optional[ExportStyleConfig] = None,
    ) -> bytes:
        """Main entry point."""
        data = self._coerce_json(payload)
        cfg = config or ExportStyleConfig()

        doc = Document()
        self._setup_numbering(doc)
        self._declare_named_styles(doc)
        self._setup_footer(doc, cfg)

        # ----- Title & Date -------------------------------------------
        title_para = doc.add_heading(data.get("title", "Contract Document"), 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_para = doc.add_paragraph(data.get("date", datetime.now().strftime("%B %d, %Y")))
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # ----- General Sections --------------------------------------
        self._write_general_sections(doc, data.get("general_sections", {}))

        # Separate clauses visually
        doc.add_page_break()

        # ----- Clauses ------------------------------------------------
        self._write_clauses(doc, data.get("clauses", []), cfg)

        # ----- Output -------------------------------------------------
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    # ==================================================================
    # JSON handling
    # ==================================================================
    @staticmethod
    def _coerce_json(src: Union[str, Dict[str, Any]]) -> Dict[str, Any]:
        if isinstance(src, str):
            return json.loads(src)
        if isinstance(src, dict):
            return src
        raise TypeError("Payload must be dict or JSON string")

    # ==================================================================
    # NUMBERING SETUP (Simplified and Fixed)
    # ==================================================================
    def _setup_numbering(self, doc: Document):
        """Setup comprehensive native DOCX numbering using XML - Simplified version"""
        try:
            # Check if numbering part already exists
            numbering_part = doc.part.numbering_part
            if numbering_part is not None:
                self.numbering_part = numbering_part
                return
        except (AttributeError, KeyError):
            pass

        # Create numbering part with proper multi-level numbering
        numbering_part_xml = f'''<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<w:numbering {nsdecls('w')}>
    <w:abstractNum w:abstractNumId="0">
        <w:nsid w:val="FFFFFFFF"/>
        <w:multiLevelType w:val="multilevel"/>
        <w:tmpl w:val="FFFFFFFF"/>
        <w:lvl w:ilvl="0">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="432" w:hanging="432"/>
            </w:pPr>
            <w:rPr>
                <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                <w:sz w:val="22"/>
            </w:rPr>
        </w:lvl>
        <w:lvl w:ilvl="1">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1.%2."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="864" w:hanging="432"/>
            </w:pPr>
            <w:rPr>
                <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                <w:sz w:val="22"/>
            </w:rPr>
        </w:lvl>
        <w:lvl w:ilvl="2">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1.%2.%3."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="1296" w:hanging="432"/>
            </w:pPr>
            <w:rPr>
                <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                <w:sz w:val="22"/>
            </w:rPr>
        </w:lvl>
        <w:lvl w:ilvl="3">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1.%2.%3.%4."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="1728" w:hanging="432"/>
            </w:pPr>
            <w:rPr>
                <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                <w:sz w:val="22"/>
            </w:rPr>
        </w:lvl>
        <w:lvl w:ilvl="4">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1.%2.%3.%4.%5."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="2160" w:hanging="432"/>
            </w:pPr>
            <w:rPr>
                <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                <w:sz w:val="22"/>
            </w:rPr>
        </w:lvl>
        <w:lvl w:ilvl="5">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1.%2.%3.%4.%5.%6."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="2592" w:hanging="432"/>
            </w:pPr>
            <w:rPr>
                <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                <w:sz w:val="22"/>
            </w:rPr>
        </w:lvl>
        <w:lvl w:ilvl="6">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1.%2.%3.%4.%5.%6.%7."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="3024" w:hanging="432"/>
            </w:pPr>
            <w:rPr>
                <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                <w:sz w:val="22"/>
            </w:rPr>
        </w:lvl>
        <w:lvl w:ilvl="7">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1.%2.%3.%4.%5.%6.%7.%8."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="3456" w:hanging="432"/>
            </w:pPr>
            <w:rPr>
                <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                <w:sz w:val="22"/>
            </w:rPr>
        </w:lvl>
        <w:lvl w:ilvl="8">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1.%2.%3.%4.%5.%6.%7.%8.%9."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="3888" w:hanging="432"/>
            </w:pPr>
            <w:rPr>
                <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                <w:sz w:val="22"/>
            </w:rPr>
        </w:lvl>
    </w:abstractNum>
    <w:num w:numId="1">
        <w:abstractNumId w:val="0"/>
    </w:num>
</w:numbering>'''

        try:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            from docx.parts.numbering import NumberingPart

            numbering_part = NumberingPart.new()
            numbering_part._element = parse_xml(numbering_part_xml)
            doc.part.package.relate_to(numbering_part, RT.NUMBERING)
            self.numbering_part = numbering_part
            self.logger.info("Successfully created numbering part")
        except Exception as e:
            self.logger.error(f"Failed to setup numbering: {e}")
            # Fall back to simple numbering without XML manipulation
            self.numbering_part = None

    # ==================================================================
    # STYLES & FOOTER
    # ==================================================================
    def get_or_make_style(self, styles, name, style_type):
        """Get existing style or create new one"""
        try:
            return styles[name]
        except KeyError:
            return styles.add_style(name, style_type)

    def _declare_named_styles(self, doc: Document):
        """Declare all named styles used in the document"""
        styles = doc.styles
        
        # Placeholder style
        ph_style = self.get_or_make_style(styles, "Placeholder", WD_STYLE_TYPE.CHARACTER)
        ph_style.font.color.rgb = RGBColor(0, 0, 0)
        ph_style.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        ph_style.font.bold = True
        
        # Drafting note paragraph style
        dn_p_style = self.get_or_make_style(styles, "DraftingNote", WD_STYLE_TYPE.PARAGRAPH)
        dn_p_style.font.color.rgb = RGBColor(25, 118, 210)
        dn_p_style.font.italic = True
        dn_p_style.paragraph_format.space_before = Pt(6)
        dn_p_style.paragraph_format.space_after = Pt(6)
        dn_p_style.paragraph_format.left_indent = Pt(36)
        
        # Drafting note character style
        dn_c_style = self.get_or_make_style(styles, "DraftingNoteChar", WD_STYLE_TYPE.CHARACTER)
        dn_c_style.font.color.rgb = RGBColor(25, 118, 210)
        dn_c_style.font.italic = True
        
        # Clause title style
        ct_style = self.get_or_make_style(styles, "ClauseTitle", WD_STYLE_TYPE.PARAGRAPH)
        ct_style.font.bold = True
        ct_style.font.size = Pt(12)
        ct_style.paragraph_format.space_before = Pt(12)
        ct_style.paragraph_format.space_after = Pt(6)
        
        # Clause body style
        cb_style = self.get_or_make_style(styles, "ClauseBody", WD_STYLE_TYPE.PARAGRAPH)
        cb_style.font.size = Pt(11)
        cb_style.paragraph_format.space_after = Pt(6)

    @staticmethod
    def _setup_footer(doc: Document, cfg: ExportStyleConfig):
        """Setup document footer with page numbers and optional logo"""
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page number field
        fld_xml = f'''<w:fldSimple {nsdecls('w')} w:instr="PAGE"><w:t>1</w:t></w:fldSimple>'''
        footer_para._p.append(parse_xml(fld_xml))
        footer_para.add_run(" of ")
        
        # Add total pages field
        fld_xml2 = f'''<w:fldSimple {nsdecls('w')} w:instr="NUMPAGES"><w:t>1</w:t></w:fldSimple>'''
        footer_para._p.append(parse_xml(fld_xml2))
        
        # Add logo if specified
        if cfg.logo_path:
            try:
                footer_para.add_run().add_picture(cfg.logo_path, height=Pt(12))
            except Exception as e:
                logger.warning(f"Failed to add logo: {e}")

    # ==================================================================
    # GENERAL SECTIONS
    # ==================================================================
    def _write_general_sections(self, doc: Document, sections: Dict[str, str]):
        """Write general sections in a specific order"""
        # Define order for important sections
        first_sections = ["preamble", "recitals", "definitions"]
        last_section = "signature_block"
        
        # Write first sections in order
        for section_key in first_sections:
            if section_key in sections:
                self._add_general_section(doc, section_key, sections[section_key])
        
        # Write remaining sections (except last)
        for key, value in sections.items():
            if key not in first_sections + [last_section]:
                self._add_general_section(doc, key, value)
        
        # Write last section
        if last_section in sections:
            self._add_general_section(doc, last_section, sections[last_section])

    def _add_general_section(self, doc: Document, key: str, html: str):
        """Add a general section with proper formatting"""
        # Add section heading
        doc.add_heading(key.replace("_", " ").title(), level=1)
        
        # Clean HTML and add content
        cleaned_text = self._strip_html(html)
        for line in cleaned_text.split("\n"):
            if line.strip():
                para = doc.add_paragraph()
                self._insert_placeholder_runs(para, line.strip())
        
        # Add spacing
        doc.add_paragraph()

    @staticmethod
    def _strip_html(txt: str) -> str:
        """Strip HTML tags while preserving line breaks"""
        # Convert line break tags to newlines
        txt = re.sub(r"<\s*br\s*/?>", "\n", txt, flags=re.I)
        txt = re.sub(r"</p>", "\n", txt, flags=re.I)
        txt = re.sub(r"<h[1-6][^>]*>", "\n", txt, flags=re.I)
        txt = re.sub(r"</h[1-6]>", "\n", txt, flags=re.I)
        
        # Remove all other HTML tags
        return re.sub(r"<[^>]+>", "", txt)

    # ==================================================================
    # CLAUSES & SUBCLAUSES WITH IMPROVED NUMBERING
    # ==================================================================
    def _write_clauses(self, doc: Document, clauses: list, cfg: ExportStyleConfig):
        """Write clauses with proper numbering and formatting"""
        if not clauses:
            return
        
        self.logger.info(f"Writing {len(clauses)} clauses")
        self._add_clause_items(doc, clauses, cfg=cfg, level=0)

    def _add_clause_items(
        self, doc: Document, items: list, cfg: ExportStyleConfig, level: int = 0
    ):
        """Add clause items with recursive nesting support"""
        if level > 8:  # Increased from 7 to 8 for more nesting levels
            self.logger.warning(f"Maximum nesting level exceeded (level {level})")
            raise ValueError("More than 9 nesting levels not supported")

        for idx, item in enumerate(items, 1):
            try:
                self._process_single_clause_item(doc, item, cfg, level)
            except Exception as e:
                self.logger.error(f"Error processing clause item {idx} at level {level}: {e}")
                # Continue with next item instead of failing completely
                continue

    def _process_single_clause_item(self, doc: Document, item: Dict, cfg: ExportStyleConfig, level: int):
        """Process a single clause item with proper formatting"""
        # Create main paragraph for the clause
        para = doc.add_paragraph()
        self._apply_numbering(para, level)

        # Handle clause title and text
        if "clause_title" in item:
            title_text = item['clause_title']
            self._insert_placeholder_runs(para, title_text)
            
            # Make title bold (except placeholders)
            if para.runs:
                for run in para.runs:
                    if not (run.text.startswith("[") and run.text.endswith("]")):
                        run.bold = True
            
            # Add clause body if present
            if clause_text := item.get("clause_text"):
                body_para = doc.add_paragraph(style="ClauseBody")
                body_para.paragraph_format.left_indent = Pt((level + 1) * 36)
                self._insert_placeholder_runs(body_para, clause_text)
                
        elif "subclause_text" in item:
            # Handle subclauses
            self._insert_placeholder_runs(para, item.get('subclause_text', ''))

        # Handle drafting notes
        if note := item.get("drafting_note"):
            if cfg.drafting_note_as_comment:
                self._add_comment(doc, para, note)
            else:
                self._add_inline_drafting_note(doc, note, level)

        # Handle nested subclauses recursively
        if subclauses := item.get("subclauses", []):
            self.logger.debug(f"Processing {len(subclauses)} subclauses at level {level + 1}")
            self._add_clause_items(doc, subclauses, cfg, level + 1)

    def _add_inline_drafting_note(self, doc: Document, note: str, level: int):
        """Add drafting note as inline paragraph"""
        note_para = doc.add_paragraph(style="DraftingNote")
        note_para.paragraph_format.left_indent = Pt(((level + 1) * 36) + 18)
        note_para.add_run(f"📝 Drafting Note: {note}")

    def _apply_numbering(self, paragraph, level: int):
        """Apply numbering to paragraph at specified level"""
        if self.numbering_part is None:
            # Fallback: add manual numbering if XML numbering failed
            self._apply_manual_numbering(paragraph, level)
            return
            
        try:
            # Create numbering properties element
            num_pr = OxmlElement('w:numPr')
            
            # Set the level
            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), str(level))
            num_pr.append(ilvl)
            
            # Set the numbering ID
            num_id_elem = OxmlElement('w:numId')
            num_id_elem.set(qn('w:val'), str(self.num_id))
            num_pr.append(num_id_elem)
            
            # Add to paragraph properties
            if paragraph._p.pPr is None:
                paragraph._p._add_pPr()
            paragraph._p.pPr.append(num_pr)
            
        except Exception as e:
            self.logger.warning(f"Failed to apply XML numbering: {e}, falling back to manual")
            self._apply_manual_numbering(paragraph, level)

    def _apply_manual_numbering(self, paragraph, level: int):
        """Fallback manual numbering when XML numbering fails"""
        indent_level = level * 0.5  # 0.5 inch per level
        paragraph.paragraph_format.left_indent = Inches(indent_level)
        # Note: This won't create actual auto-numbering, just indentation

    # ==================================================================
    # PLACEHOLDER AND FORMATTING
    # ==================================================================
    def _insert_placeholder_runs(self, paragraph, text: str):
        """Insert text with placeholder formatting"""
        if not text:
            return
            
        # Split text by placeholders while preserving the delimiters
        parts = self.PLACEHOLDER_REGEX.split(text)
        
        for part in parts:
            if not part:
                continue
                
            if part.startswith("[") and part.endswith("]"):
                # This is a placeholder
                run = paragraph.add_run(part)
                self._format_placeholder_run(run)
            else:
                # Regular text
                paragraph.add_run(part)

    def _format_placeholder_run(self, run):
        """Format a placeholder run with highlighting"""
        try:
            # Ensure run properties exist
            if run._r.rPr is None:
                run._r._add_rPr()
            
            rPr = run._r.rPr
            
            # Set color to black
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '000000')
            rPr.append(color)
            
            # Set highlight to green
            highlight = OxmlElement('w:highlight')
            highlight.set(qn('w:val'), 'green')
            rPr.append(highlight)
            
            # Set bold
            bold = OxmlElement('w:b')
            rPr.append(bold)
            
        except Exception as e:
            self.logger.warning(f"Failed to format placeholder run: {e}")
            # Fallback to simple formatting
            run.bold = True

    # ==================================================================
    # COMMENTS (Fixed Implementation)
    # ==================================================================
    def _add_comment(self, doc: Document, paragraph, text: str):
        """Add comment using native DOCX features"""
        try:
            self._add_native_comment(doc, paragraph, text, author="DocGen", initials="DG")
        except Exception as e:
            self.logger.warning(f"Failed to add native comment: {e}, falling back to inline")
            self._add_fallback_comment(paragraph, text)

    def _add_native_comment(self, doc: Document, paragraph, text: str, author: str, initials: str):
        """
        Add a native DOCX comment using the correct python-docx API
        Fixed: Proper parameter order and usage
        """
        try:
            # Ensure paragraph has content to anchor comment to
            if not paragraph.runs:
                # Add an empty run if paragraph is empty
                paragraph.add_run("")
            
            # Add comment using correct API signature
            comment = doc.add_comment(
                text=text,
                author=author,
                initials=initials,
                runs=paragraph.runs  # Anchor to all runs in the paragraph
            )
            
            self.comment_counter += 1
            self.logger.info(f"Added comment #{self.comment_counter}: '{text[:50]}...'")
            return comment
            
        except Exception as e:
            self.logger.error(f"Native comment creation failed: {e}")
            raise

    def _add_fallback_comment(self, paragraph, text: str):
        """Fallback: Add comment as inline formatted text"""
        comment_run = paragraph.add_run(f" [Comment: {text}]")
        self._format_comment_run(comment_run)

    def _format_comment_run(self, run):
        """Format fallback comment run"""
        try:
            if run._r.rPr is None:
                run._r._add_rPr()
            
            rPr = run._r.rPr
            
            # Set color to blue
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '1976D2')
            rPr.append(color)
            
            # Set italic
            italic = OxmlElement('w:i')
            rPr.append(italic)
            
            # Set smaller font size
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '20')  # 10pt font
            rPr.append(sz)
            
        except Exception as e:
            self.logger.warning(f"Failed to format comment run: {e}")
            # Simple fallback
            run.italic = True

    # ==================================================================
    # UTILITY METHODS
    # ==================================================================
    def get_stats(self) -> Dict[str, int]:
        """Get export statistics"""
        return {
            "comments_added": self.comment_counter,
            "numbering_part_created": self.numbering_part is not None
        }
